from docx import Document
from docx.shared import Inches
import streamlit as st


def build_report_view(summaries, articles):
    st.markdown("### 🗂️ Rapport structuré par sujet")
    for topic, summary in summaries.items():
        st.markdown(f"#### 🧠 {topic}")
        st.markdown(summary)

    st.markdown("### 🔗 Sources utilisées")
    for article in articles:
        st.markdown(f"- [{article['title']}]({article['link']})")


def generate_docx(summaries, articles):
    doc = Document()
    doc.add_heading("Rapport de veille stratégique IA", 0)

    doc.add_heading("Synthèses par sujet", level=1)
    for topic, summary in summaries.items():
        doc.add_heading(topic, level=2)
        doc.add_paragraph(summary)

    doc.add_page_break()
    doc.add_heading("Sources consultées", level=1)
    for article in articles:
        doc.add_paragraph(f"- {article['title']} ({article['link']})")

    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

