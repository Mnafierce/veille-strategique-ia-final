from io import BytesIO
from docx import Document

def generate_docx(summaries, articles):
    doc = Document()
    doc.add_heading("Rapport de veille stratégique – Agents IA", 0)

    for topic, summary in summaries.items():
        doc.add_heading(topic, level=1)
        doc.add_paragraph(summary)

        # Ajouter les sources associées
        related = [a for a in articles if a["keyword"] == topic]
        if related:
            doc.add_paragraph("Sources :", style="Intense Quote")
            for a in related:
                doc.add_paragraph(f"- {a['title']} ({a['link']})", style="List Bullet")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
